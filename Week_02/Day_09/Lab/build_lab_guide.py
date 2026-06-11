r"""Build Day_09_Lab_Guide.docx from Lab_Files/LABS_GUIDE.md.

A teacher-friendly, student-facing GUIDE — every command is shown, every
command is explained, every step says exactly what to do and why.

Run from Day_09/ root:
    .\Slides\.venv\Scripts\python.exe Lab\build_lab_guide.py
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).parent
OUT = HERE / "Day_09_Lab_Guide.docx"


# -------------------------------- styling --------------------------------- #
ACCENT = RGBColor(0x0E, 0x9F, 0x8F)     # teal
HEADING = RGBColor(0x14, 0x18, 0x20)
MUTED = RGBColor(0x4A, 0x52, 0x60)
INFO_BG = "EAF6F4"
WARN_BG = "FFF5E6"
CODE_BG = "F3F5F8"
CODE_FG = RGBColor(0x14, 0x18, 0x20)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=120, right=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def add_code_block(doc, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, CODE_BG)
    set_cell_margins(cell)
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = CODE_FG


def add_callout(doc, label: str, body: str,
                bg: str = INFO_BG, label_color=ACCENT) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    p = cell.paragraphs[0]
    lab = p.add_run(f"{label}  ")
    lab.bold = True
    lab.font.color.rgb = label_color
    lab.font.size = Pt(10)
    body_run = p.add_run(body)
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = HEADING


def add_heading(doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADING
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)


def add_para(doc, text: str, *, italic: bool = False, bold: bool = False,
             color=None, size: int = 11) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color


def add_bullets(doc, items) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)


def add_section_label(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def add_step(doc, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(11)


def add_why(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"What this does:  {text}")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED


def add_lab_block(doc, *, num: int, title: str, goal: str,
                  context: str, steps: list,
                  observe: list, discuss: list,
                  tip: str | None = None,
                  warn: str | None = None) -> None:
    """Render one lab.

    Each step is a tuple. Supported kinds:
        ("step",  "1. Do something")          -> numbered instruction
        ("why",   "explains the command")     -> italic explanation paragraph
        ("code",  "command-line text")        -> shaded mono block
        ("note",  "callout-style text")       -> teal callout
        ("warn",  "callout-style text")       -> amber callout
    """
    add_heading(doc, f"Lab {num} — {title}", level=2)

    add_section_label(doc, "Goal")
    add_para(doc, goal, size=11)

    add_section_label(doc, "Why this lab matters")
    add_para(doc, context, italic=True, color=MUTED, size=10)

    if tip:
        add_callout(doc, "TIP:", tip, bg=INFO_BG)
    if warn:
        add_callout(doc, "READ THIS FIRST:", warn,
                    bg=WARN_BG, label_color=RGBColor(0xB8, 0x6E, 0x00))

    add_section_label(doc, "Step-by-step")
    for kind, payload in steps:
        if kind == "step":
            add_step(doc, payload)
        elif kind == "why":
            add_why(doc, payload)
        elif kind == "code":
            add_code_block(doc, payload)
        elif kind == "note":
            add_callout(doc, "TIP:", payload, bg=INFO_BG)
        elif kind == "warn":
            add_callout(doc, "WATCH OUT:", payload, bg=WARN_BG,
                        label_color=RGBColor(0xB8, 0x6E, 0x00))

    add_section_label(doc, "What to observe")
    add_bullets(doc, observe)

    add_section_label(doc, "Discussion questions")
    add_bullets(doc, discuss)

    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(0)
    sep.add_run("―" * 60).font.color.rgb = MUTED


# ============================================================
def build() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    title = doc.add_paragraph()
    r = title.add_run("Day 09 — Student Lab Guide")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = HEADING

    sub = doc.add_paragraph()
    rs = sub.add_run(
        "Pipeline Monitoring · Logging & Error Handling · "
        "Data Quality · Scaling"
    )
    rs.italic = True
    rs.font.size = Pt(14)
    rs.font.color.rgb = ACCENT

    doc.add_paragraph(
        "This guide walks you through 15 hands-on labs on a real NYC-taxi "
        "streaming pipeline. Every command you need to type is shown in a "
        "grey box. Underneath each command we explain WHAT IT DOES — read "
        "those lines, they are the actual learning.\n\n"
        "Plan around 5 hours total. Each lab takes ~20 minutes. You can "
        "stop and resume after any lab as long as the Docker stack stays "
        "running."
    )

    add_callout(doc, "Stack:",
        "KRaft Kafka (3 brokers) · Postgres + Debezium CDC · ksqlDB · "
        "Schema Registry · Prometheus · Grafana · Loki · Alertmanager · "
        "Kafka UI · expectation-based validator")
    add_callout(doc, "Dataset:",
        "NYC Uber-style taxi data — lat/lon, fares, surge multipliers, "
        "payment types, and 100 drivers sourced from Postgres via CDC.")
    add_callout(doc, "You will need:",
        "Docker Desktop (Windows/macOS) or Docker Engine + Compose plug-in "
        "(Linux/WSL), Python 3.10+, ~6 GB free RAM, ~8 GB disk space, and "
        "ports 3000, 5000, 8080, 8081, 8083, 8088, 9090, 9092-9096 free.",
        bg=WARN_BG, label_color=RGBColor(0xB8, 0x6E, 0x00))

    # ============================================================
    # HOW TO READ THIS GUIDE
    # ============================================================
    add_heading(doc, "How to Read This Guide", level=1)
    add_bullets(doc, [
        "Lines in a grey monospace box are commands. Type them exactly as "
            "shown. On Linux/macOS/WSL/Git Bash use the bash version; on "
            "Windows cmd use the cmd version where two are given.",
        "Italic 'What this does' lines explain WHY you're running the "
            "command. They are the actual lesson — read them, don't skip.",
        "Teal call-outs are TIPS that save you time.",
        "Amber call-outs are warnings — read them BEFORE the next command.",
        "Each lab ends with 'What to observe' and 'Discussion questions'. "
            "Work through both before moving on. Discussion questions are "
            "interview questions in disguise.",
    ])

    # ============================================================
    # ARCHITECTURE
    # ============================================================
    add_heading(doc, "Architecture at a Glance", level=1)
    add_para(doc,
        "This is the system you will be operating today. Every box is a "
        "container; every arrow is a Kafka topic.")

    add_code_block(doc,
        "  Postgres.drivers  ──►  Debezium  ──►  cdc.public.drivers (CDC topic, KTable)\n"
        "                                                  │\n"
        "                                                  ▼\n"
        "  taxi_simulator.py ─► gps-pings  (12p, RF=3)     │\n"
        "                       taxi-trips (6p,  RF=3) ─►  taxi_consumer\n"
        "                                                   │           │\n"
        "                                                   ▼           ▼\n"
        "                                          trips-clean  ─► driver_enricher ─► trips-enriched\n"
        "                                            │      │                              │\n"
        "                                            │      ▼                              │\n"
        "                                            │  surge_detector ─► surge-events     │\n"
        "                                            │  quality_validator                  │\n"
        "                                            ▼                                     ▼\n"
        "                                        trips-dlq  ◄────── dashboard.py ◄─────  │\n"
        "                                                            (live map @ :5000)\n\n"
        "  Observability:  JMX → Prometheus → Grafana    logs → Promtail → Loki → Grafana"
    )

    add_para(doc, "Medallion-architecture mapping (bronze/silver/gold):",
             bold=True, size=11)
    tbl = doc.add_table(rows=5, cols=3)
    tbl.style = "Light List Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Layer", "Topic", "What lives there"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    rows = [
        ("Bronze",     "taxi-trips",      "Raw simulator events (~1.5 % bad data)"),
        ("Silver",     "trips-clean",     "Schema-validated, light checks passed"),
        ("Gold",       "trips-enriched",  "Joined with driver KTable from Postgres CDC"),
        ("Quarantine", "trips-dlq",       "Rejected records with _dlq_reason"),
    ]
    for i, (layer, topic, mean) in enumerate(rows, start=1):
        c = tbl.rows[i].cells
        c[0].text = layer
        c[1].text = topic
        c[2].text = mean

    # ============================================================
    # PREREQUISITES — DETAILED
    # ============================================================
    doc.add_page_break()
    add_heading(doc, "Before You Start", level=1)
    add_para(doc,
        "You have two paths: the FAST PATH (use the bootstrap script that "
        "does everything for you) or the MANUAL PATH (run the same steps "
        "yourself so you understand each one). Pick one. We recommend the "
        "fast path the first time, then re-do the manual path later if you "
        "want to learn the plumbing.")

    # ---- Fast path ----
    add_heading(doc, "Fast Path — One Command Sets Everything Up", level=2)
    add_para(doc, "Open a terminal in the Day_09/Lab_Files folder.")

    add_section_label(doc, "Linux · macOS · WSL · Git Bash")
    add_code_block(doc,
        "cd Lab_Files\n"
        "chmod +x bootstrap.sh\n"
        "./bootstrap.sh")
    add_why(doc,
        "downloads the JMX javaagent, creates a Python virtualenv at "
        ".venv/, installs requirements.txt, runs 'docker compose up -d' "
        "to start all 15 containers, calls setup_topics.py to create "
        "every topic with the right partitions/replication, seeds 100 "
        "driver rows into Postgres, and registers the Debezium "
        "PostgreSQL connector.")

    add_section_label(doc, "Windows (cmd / PowerShell)")
    add_code_block(doc,
        "cd Lab_Files\n"
        "bootstrap.cmd")
    add_why(doc, "same as above, using the Windows batch version.")

    add_callout(doc, "First time only:",
        "The bootstrap may take 3-5 minutes the first time because Docker "
        "has to pull about 4 GB of images. Subsequent runs are seconds.")

    # ---- Manual path ----
    add_heading(doc, "Manual Path — If You Want to See Every Step", level=2)
    add_para(doc,
        "Run these in order. Each block has a 'What this does' line so "
        "you can follow along.")

    add_section_label(doc, "1) Create the Python virtual environment")
    add_code_block(doc,
        "# Linux / macOS / WSL / Git Bash\n"
        "cd Lab_Files\n"
        "python3 -m venv .venv\n"
        "source .venv/bin/activate\n"
        "pip install --upgrade pip\n"
        "pip install -r requirements.txt")
    add_code_block(doc,
        ":: Windows cmd\n"
        "cd Lab_Files\n"
        "python -m venv .venv\n"
        ".venv\\Scripts\\activate.bat\n"
        "python -m pip install --upgrade pip\n"
        "pip install -r requirements.txt")
    add_why(doc,
        "installs kafka-python, confluent-kafka, prometheus-client, "
        "flask and a few other libs into a project-local Python so it "
        "cannot break your system Python.")

    add_section_label(doc, "2) Download the JMX Prometheus javaagent")
    add_code_block(doc,
        "# Linux / macOS / WSL / Git Bash\n"
        "mkdir -p jmx_exporter\n"
        "curl -L -o jmx_exporter/jmx_prometheus_javaagent.jar \\\n"
        "  https://repo1.maven.org/maven2/io/prometheus/jmx/jmx_prometheus_javaagent/0.20.0/jmx_prometheus_javaagent-0.20.0.jar")
    add_why(doc,
        "fetches the JAR that bridges Kafka's JMX MBeans to a /metrics "
        "HTTP endpoint Prometheus can scrape. Without it, Prometheus "
        "has nothing to read from the brokers.")

    add_section_label(doc, "3) Start the whole Docker stack")
    add_code_block(doc, "docker compose up -d")
    add_why(doc,
        "reads docker-compose.yml and starts ~15 containers in the "
        "background ('-d' = detached). The first time this pulls ~4 GB "
        "of images. The brokers take ~30 seconds to elect a controller, "
        "so wait until 'docker compose ps' shows them all 'running' "
        "before continuing.")
    add_code_block(doc, "docker compose ps")
    add_why(doc,
        "shows the state of every container so you can verify all 15 "
        "are running.")

    add_section_label(doc, "4) Create the Kafka topics")
    add_code_block(doc, "python setup_topics.py")
    add_why(doc,
        "creates taxi-trips (6 partitions, RF=3), gps-pings (12, RF=3), "
        "trips-clean, trips-enriched, trips-dlq (24 h retention) and "
        "surge-events (compacted). Idempotent — safe to run more than "
        "once.")

    add_section_label(doc, "5) Seed Postgres with driver data")
    add_code_block(doc, "python db_seeder.py")
    add_why(doc,
        "inserts 100 rows into the drivers table inside the postgres "
        "container. Debezium will then publish these via CDC.")

    add_section_label(doc, "6) Register the Debezium connector")
    add_code_block(doc, "python register_connector.py")
    add_why(doc,
        "POSTs debezium-postgres.json to Kafka Connect, which starts "
        "streaming Postgres write-ahead-log entries into the "
        "cdc.public.drivers topic.")

    # ---- Activate in every new terminal ----
    add_heading(doc, "Every Time You Open a New Terminal", level=2)
    add_para(doc,
        "Most labs need 3-5 terminals open at once. In EVERY new terminal, "
        "first navigate into Lab_Files/ and activate the venv so 'python' "
        "points at the right Python with the right packages.")

    add_code_block(doc,
        "# Linux / macOS / WSL / Git Bash\n"
        "cd Lab_Files\n"
        "source .venv/bin/activate")
    add_code_block(doc,
        ":: Windows cmd\n"
        "cd Lab_Files\n"
        ".venv\\Scripts\\activate.bat")
    add_code_block(doc,
        "# Windows PowerShell\n"
        "cd Lab_Files\n"
        "Set-ExecutionPolicy -Scope Process -ExecutionPolicy RemoteSigned\n"
        ".\\.venv\\Scripts\\Activate.ps1")
    add_why(doc,
        "prepends the venv's Scripts/ folder to PATH so the 'python' you "
        "call is the one with all the dependencies installed.")

    # ---- Smoke test ----
    add_heading(doc, "Smoke-Test Everything Is Up", level=2)
    add_code_block(doc,
        "docker compose ps                            # 15 containers, all 'running'\n"
        "curl http://localhost:8080                   # Kafka UI HTML\n"
        "curl http://localhost:9090/-/ready           # 'Prometheus is Ready.'\n"
        "curl http://localhost:3000                   # Grafana HTML\n"
        "curl http://localhost:8083/connectors        # [\"drivers-postgres-source\"]")
    add_why(doc,
        "verifies each piece of the stack is reachable on its expected "
        "host port before you start the labs. If any of these fail, fix "
        "the failing one first (most often it's still starting — wait "
        "30 seconds and retry).")

    add_callout(doc, "Tear down at the end:",
        "When the whole session is over, run 'docker compose down -v' "
        "from Lab_Files/ to stop and remove every container plus their "
        "data volumes.")

    # ============================================================
    # PART 1 — PIPELINE & DATA SOURCES
    # ============================================================
    doc.add_page_break()
    add_heading(doc, "Part 1 · Pipeline & Data Sources (Labs 1-4)", level=1)
    add_para(doc,
        "First we explore the cluster, then we wire CDC into the pipeline "
        "and finish with the live map dashboard — the 'wow moment' of "
        "the day.", italic=True, color=MUTED)

    # ------------------ LAB 1 ------------------
    add_lab_block(doc, num=1,
        title="KRaft Cluster & Topic Design",
        goal=("Confirm the cluster is up, understand why there is no "
              "Zookeeper, and read the partition/replication design of "
              "every topic."),
        context=("Every Kafka topic you'll touch today is created with a "
                 "deliberate partition count and retention policy. If you "
                 "don't know the layout, you can't reason about lag, "
                 "throughput or compaction. We start here for that reason."),
        tip=("All 'docker exec' commands run INSIDE the kafka1 container "
             "so they can talk to brokers on the internal port 29092. "
             "From your laptop you'd use localhost:9094 instead."),
        steps=[
            ("step", "1. Confirm the Docker stack is running. If not, start it now."),
            ("code", "docker compose ps"),
            ("why",  "lists every container with its state. You should see 15 rows, all 'running' or 'healthy'."),
            ("code", "docker compose up -d        # run only if any container is missing"),
            ("why",  "'-d' starts everything in the background. Idempotent — safe to re-run."),

            ("step", "2. Check the KRaft quorum status (no Zookeeper container exists)."),
            ("code",
                "docker exec kafka1 kafka-metadata-quorum \\\n"
                "  --bootstrap-server kafka1:29092 describe --status"),
            ("why",  "prints the current controller, the leader epoch and the highest committed offset. KRaft replaces Zookeeper — the quorum lives inside the brokers themselves."),

            ("step", "3. List every topic the bootstrap created."),
            ("code",
                "docker exec kafka1 kafka-topics --bootstrap-server kafka1:29092 --list"),
            ("why",  "you should see taxi-trips, gps-pings, trips-clean, trips-enriched, trips-dlq, surge-events, cdc.public.drivers and a few internal __ topics."),

            ("step", "4. Describe each topic and read its partition count, replication and configs."),
            ("code",
                "docker exec kafka1 kafka-topics --bootstrap-server kafka1:29092 --describe --topic taxi-trips\n"
                "docker exec kafka1 kafka-topics --bootstrap-server kafka1:29092 --describe --topic gps-pings\n"
                "docker exec kafka1 kafka-topics --bootstrap-server kafka1:29092 --describe --topic surge-events\n"
                "docker exec kafka1 kafka-topics --bootstrap-server kafka1:29092 --describe --topic trips-dlq"),
            ("why",  "reads the topic config from Kafka itself — partition count, replicas, ISR, min.insync.replicas, cleanup.policy and retention."),
        ],
        observe=[
            "There is no Zookeeper container — KRaft quorum is hosted by the three Kafka processes.",
            "taxi-trips has 6 partitions, gps-pings has 12 (10× higher volume justifies more parallelism).",
            "surge-events uses cleanup.policy=compact — only the LATEST surge value per zone is kept (compacted topic ≈ KTable).",
            "trips-dlq has retention.ms set to 24 h so you have time to triage failures before they expire.",
        ],
        discuss=[
            "Why partition gps-pings more aggressively than taxi-trips?",
            "Why is cleanup.policy=compact appropriate for surge-events but NOT for taxi-trips?",
            "What does min.insync.replicas=2 guarantee on a 3-broker cluster?",
        ],
    )

    # ------------------ LAB 2 ------------------
    add_lab_block(doc, num=2,
        title="Cluster Inspection with Kafka UI",
        goal="Get fluent with the operator's GUI before going terminal-only.",
        context=("Kafka UI is the same kind of tool real operations teams "
                 "use on Confluent Cloud, AKHQ or Redpanda Console. "
                 "Learning to navigate it is a transferable production "
                 "skill."),
        steps=[
            ("step", "1. Open Kafka UI in your browser."),
            ("code", "http://localhost:8080"),
            ("why",  "pick the 'taxi-kraft-cluster' from the cluster list on the left."),

            ("step", "2. Click the Brokers tab."),
            ("why",  "confirms 3 brokers and shows their JMX metrics (controller, partition count, leader count)."),

            ("step", "3. Click the Topics tab → click 'taxi-trips' → click 'Messages'."),
            ("why",  "the Messages view streams live records as they arrive. Empty for now because nothing is producing."),

            ("step", "4. Click the Schema Registry tab."),
            ("why",  "will be empty until Lab 10 — visit anyway so you know where it lives."),

            ("step", "5. Click the KSQL DB tab."),
            ("why",  "confirms ksqlDB is reachable. You will use it in Lab 12."),

            ("step", "6. Open a NEW terminal (activate the venv first) and start the simulator briefly:"),
            ("code",
                "# new terminal, in Lab_Files/, venv activated\n"
                "python taxi_simulator.py --drivers 20"),
            ("why",  "spawns 20 virtual drivers that produce taxi-trips and gps-pings records. Leave it running for ~30 seconds."),

            ("step", "7. Refresh Kafka UI → Topics → taxi-trips → Messages."),
            ("why",  "you should see records arriving in real time. Headers, keys, values and offsets are all visible."),

            ("step", "8. Stop the simulator with Ctrl+C in its terminal."),
            ("why",  "producing more data than needed will fill up disks during a long lab session."),
        ],
        observe=[
            "Per-partition message counts in 'Overview' are roughly balanced — good key distribution.",
            "Each message has a key (driver_id), a JSON value, and Kafka headers.",
            "Consumer Groups tab is EMPTY — no one is consuming yet, so messages just sit on disk.",
        ],
        discuss=[
            "Why does the simulator key messages by driver_id?",
            "What would happen if every message used the SAME key?",
        ],
    )

    # ------------------ LAB 3 ------------------
    add_lab_block(doc, num=3,
        title="Multiple Sources — Postgres + Debezium CDC",
        goal=("Combine an event stream (taxi trips) with a reference table "
              "(drivers) using Change Data Capture so you get a "
              "stream-table join end-to-end."),
        context=("Real pipelines mix EVENTS (high-volume, append-only) with "
                 "DIMENSIONS (low-volume, mutable). CDC is the standard "
                 "way to stream a database table into Kafka. You'll do it "
                 "here with Postgres → Debezium → enricher."),
        warn=("This lab needs 3 terminals open at the same time. Open "
              "them all up front and activate the venv in each before "
              "starting."),
        steps=[
            ("step", "1. Inspect the seeded drivers table in Postgres."),
            ("code",
                "docker exec -it postgres psql -U taxi -d taxi -c \\\n"
                "  \"SELECT driver_id, full_name, vehicle_make, rating FROM drivers LIMIT 5;\""),
            ("why",  "connects to the postgres container and runs a SQL query. You should see DRV-0001 through DRV-0005."),

            ("step", "2. Verify the Debezium connector is RUNNING."),
            ("code",
                "curl -s http://localhost:8083/connectors/drivers-postgres-source/status | python -m json.tool"),
            ("why",  "Kafka Connect exposes a REST API on 8083. Look for connector.state = RUNNING and tasks[0].state = RUNNING."),

            ("step", "3. Confirm the CDC topic contains a snapshot of all 100 drivers."),
            ("code",
                "docker exec kafka1 kafka-console-consumer \\\n"
                "  --bootstrap-server kafka1:29092 \\\n"
                "  --topic cdc.public.drivers \\\n"
                "  --from-beginning --max-messages 3"),
            ("why",  "reads the first 3 messages from the CDC topic. Debezium emits one record per row in the initial snapshot, then streams future changes."),

            ("step", "4. Live CDC test — change a driver in Postgres while watching Kafka."),
            ("code",
                "docker exec -it postgres psql -U taxi -d taxi -c \\\n"
                "  \"UPDATE drivers SET rating = 4.99 WHERE driver_id = 'DRV-0001';\""),
            ("why",  "within ~1 second a new message appears on cdc.public.drivers for DRV-0001 carrying the new rating value."),

            ("step", "5. Open Terminal A and start the enricher."),
            ("code",
                "# Terminal A — Lab_Files/, venv activated\n"
                "python driver_enricher.py"),
            ("why",  "materialises cdc.public.drivers into an in-memory dict, then joins every trip with its driver's name/vehicle/rating."),

            ("step", "6. Open Terminal B and start the simulator."),
            ("code",
                "# Terminal B — Lab_Files/, venv activated\n"
                "python taxi_simulator.py --drivers 30"),
            ("why",  "produces taxi-trips and gps-pings. The enricher will start emitting joined records to trips-enriched."),

            ("step", "7. Open Terminal C and start the consumer."),
            ("code",
                "# Terminal C — Lab_Files/, venv activated\n"
                "python taxi_consumer.py"),
            ("why",  "validates trips and forwards good ones to trips-clean, bad ones to trips-dlq."),

            ("step", "8. In a fourth terminal, inspect an enriched trip."),
            ("code",
                "docker exec kafka1 kafka-console-consumer \\\n"
                "  --bootstrap-server kafka1:29092 \\\n"
                "  --topic trips-enriched \\\n"
                "  --from-beginning --max-messages 1 | python -m json.tool"),
            ("why",  "the output now contains driver_name, vehicle, driver_rating and license_number alongside the original trip fields. That is the stream-table join in action."),

            ("step", "9. Stop the simulator (Ctrl+C) when you've seen enough. Leave the enricher running for Lab 4."),
        ],
        observe=[
            "CDC turns the database table into a Kafka topic, automatically maintained.",
            "driver_enricher.py implements the canonical Kafka Streams stream-table JOIN in pure Python — easy to read.",
            "A single UPDATE in Postgres flows through Debezium → Kafka → enricher in under 1 second.",
        ],
        discuss=[
            "What's the difference between Debezium's snapshot phase and streaming phase?",
            "Why does Debezium need wal_level=logical in Postgres?",
            "When would you reach for CDC vs. periodic batch sync (e.g. nightly dump)?",
        ],
    )

    # ------------------ LAB 4 ------------------
    add_lab_block(doc, num=4,
        title="The Live Map Dashboard (the wow moment)",
        goal=("See the entire pipeline running on a live NYC map, with "
              "DLQ and Data Quality panels updating in real time."),
        context=("You've seen every piece individually. Now run them all "
                 "together and watch the system come to life. This is the "
                 "shape of a real production streaming app."),
        tip=("Open all 5 terminals first, activate the venv in each, "
             "THEN start the processes in the listed order. Starting "
             "them out of order is fine but you'll see error messages "
             "until each downstream consumer finds its upstream topic."),
        steps=[
            ("step", "1. Open 5 terminals. In every one, run:"),
            ("code",
                "cd Lab_Files\n"
                "source .venv/bin/activate                # Linux / macOS / WSL\n"
                ":: OR\n"
                ".venv\\Scripts\\activate.bat              :: Windows cmd"),
            ("why",  "each script imports kafka-python and prometheus-client from the venv. Without activation you'll get ModuleNotFoundError."),

            ("step", "2. Terminal 1 — start the simulator (50 virtual drivers)."),
            ("code", "python taxi_simulator.py --drivers 50"),
            ("why",  "produces ~50 trips/s into taxi-trips, plus high-frequency GPS pings into gps-pings."),

            ("step", "3. Terminal 2 — start the bronze→silver consumer."),
            ("code", "python taxi_consumer.py"),
            ("why",  "reads taxi-trips, validates schema and basic ranges, writes good records to trips-clean and bad ones to trips-dlq with a _dlq_reason header."),

            ("step", "4. Terminal 3 — start the silver→gold enricher."),
            ("code", "python driver_enricher.py"),
            ("why",  "joins trips-clean with the driver KTable from CDC and writes the enriched output to trips-enriched."),

            ("step", "5. Terminal 4 — start the surge detector."),
            ("code", "python surge_detector.py"),
            ("why",  "computes a 30-second tumbling-window aggregation per pickup zone and writes the current surge multiplier to surge-events (a compacted topic)."),

            ("step", "6. Terminal 5 — start the quality validator."),
            ("code", "python quality_validator.py"),
            ("why",  "runs a declarative expectation suite on trips-clean and exports per-rule pass/fail counters to Prometheus."),

            ("step", "7. Open the live dashboard in your browser."),
            ("code", "http://localhost:5000"),
            ("why",  "dashboard.py (already running inside the taxi-dashboard container) subscribes to taxi-trips, surge-events and trips-dlq and pushes updates to your browser over WebSocket."),

            ("step", "8. Watch the map for 2-3 minutes."),
            ("why",  "you should see taxi emojis moving, zone circles changing colour during surges, the DLQ panel ticking up and the Quality panel showing pass/fail counters per rule."),

            ("step", "9. Leave everything running — Labs 5-12 build on this state."),
        ],
        observe=[
            "~50 taxi emojis move across Manhattan; ON_TRIP drivers are bright, IDLE drivers are dim.",
            "Zone circles change colour when the surge multiplier rises.",
            "Right-side DLQ panel shows the most recent bad records with their _dlq_reason.",
            "Right-side Quality panel shows live pass/fail counts per rule.",
            "Bottom event log streams TRIP / SURGE / DLQ lines as they happen.",
        ],
        discuss=[
            "Trace each of the four data flows on the dashboard back to its Kafka topic.",
            "What is the latency from 'trip produced' to 'appears on map'? "
            "(Open the browser DevTools → Network → WS to measure.)",
        ],
    )

    # ============================================================
    # PART 2 — MONITORING
    # ============================================================
    doc.add_page_break()
    add_heading(doc, "Part 2 · Pipeline Monitoring (Labs 5-7)", level=1)
    add_para(doc,
        "Wire up the observability stack, learn to read consumer lag, "
        "and treat logs like metrics.",
        italic=True, color=MUTED)

    # ------------------ LAB 5 ------------------
    add_lab_block(doc, num=5,
        title="JMX → Prometheus → Grafana",
        goal="Expose Kafka broker internals via JMX and visualise them in Grafana.",
        context=("Kafka brokers speak JMX (Java's built-in metrics "
                 "system). The JMX exporter is a small Java agent that "
                 "translates JMX MBeans into a Prometheus /metrics page. "
                 "Prometheus scrapes that page every 15 s and stores the "
                 "numbers. Grafana draws pictures of those numbers."),
        steps=[
            ("step", "1. Verify Prometheus is scraping every target."),
            ("code",
                "curl -s http://localhost:9090/api/v1/targets | grep -o '\"health\":\"[^\"]*\"' | sort -u"),
            ("why",  "all entries should be 'up'. If you see 'down', open the Prometheus UI and read the 'Last Error' column."),

            ("step", "2. Open Prometheus in your browser → Status → Targets."),
            ("code", "http://localhost:9090"),
            ("why",  "visually confirms each scrape job: kafka-brokers (3 targets), kafka-lag-exporter, prometheus itself, and the Python services."),

            ("step", "3. Run a few PromQL queries in the Prometheus UI (Graph tab)."),
            ("code",
                "sum(rate(kafka_server_brokertopicmetrics_messagesin_total[1m]))\n"
                "sum by (topic) (rate(kafka_server_brokertopicmetrics_bytesin_total[1m]))\n"
                "kafka_server_replicamanager_underreplicatedpartitions"),
            ("why",  "three classic queries: total ingest rate; bytes per topic per second; any partition currently under-replicated. All three should return non-empty results."),

            ("step", "4. Open Grafana in your browser."),
            ("code",
                "http://localhost:3000\n"
                "# default login:  admin / admin   (you can skip the password change)"),
            ("why",  "Grafana is the day-to-day operator view. The data source 'Prometheus' is already provisioned."),

            ("step", "5. Open the pre-built dashboard."),
            ("code", "Dashboards → Taxi Pipeline → 'Taxi Pipeline – Live Overview'"),
            ("why",  "provisioned from grafana_provisioning/dashboards/taxi_overview.json — same Trips/sec, GPS pings/sec, quality and surge panels you saw in Lab 4 but now from real Prometheus metrics."),
        ],
        observe=[
            "Per-broker 'Bytes In' should be balanced across kafka1/2/3 — uneven means a key skew problem.",
            "Bottom panels show DLQ counts and expectation failures broken out per rule.",
            "Each panel has a query — click the title → Edit to read the PromQL behind it.",
        ],
        discuss=[
            "Kafka brokers expose 1000+ JMX metrics. Why do we project only the patterns in jmx_exporter/kafka.yml?",
            "Which two metrics would you alert on first in production?",
        ],
    )

    # ------------------ LAB 6 ------------------
    add_lab_block(doc, num=6,
        title="Consumer-Group Lag Monitoring",
        goal="Lag is the #1 health signal for a Kafka consumer. Learn to read it.",
        context=("Lag = (latest offset on the partition) − (consumer's "
                 "last committed offset). Rising lag means the consumer "
                 "is falling behind. It is the EARLIEST warning that "
                 "something is wrong, almost always firing before errors "
                 "or SLAs."),
        steps=[
            ("step", "1. Check current lag from the Kafka CLI."),
            ("code",
                "docker exec kafka1 kafka-consumer-groups --bootstrap-server kafka1:29092 \\\n"
                "  --describe --group trip-processor-v1"),
            ("why",  "shows per-partition CURRENT-OFFSET, LOG-END-OFFSET and LAG for the consumer group taxi_consumer.py uses."),

            ("step", "2. Open Prometheus and confirm kafka-lag-exporter is publishing the same numbers."),
            ("code", "kafka_consumergroup_group_lag{group=\"trip-processor-v1\"}"),
            ("why",  "Prometheus has its own scraped copy of the lag, so Grafana can graph it without shelling into a container."),

            ("step", "3. Generate intentional lag — go to Terminal 2 (taxi_consumer.py) and press Ctrl+C."),
            ("why",  "with the consumer down, producer keeps writing → lag climbs."),

            ("step", "4. Watch the lag grow in Grafana for ~60 seconds."),
            ("why",  "the 'Consumer Lag' panel rises in real time. This is exactly what an SRE would see at the start of an incident."),

            ("step", "5. Restart the consumer."),
            ("code", "python taxi_consumer.py"),
            ("why",  "the consumer catches up — the lag line drains back toward zero."),
        ],
        observe=[
            "Lag spikes when the consumer is down, then drains rapidly when it returns.",
            "Each partition has its OWN lag value — uneven keys would show uneven lag.",
            "Lag is the EARLIEST warning signal — it rises minutes before SLO breaches.",
        ],
        discuss=[
            "A consumer at lag = 50 000 on a topic doing 5 000 msg/s — how far behind is it in TIME?",
            "When would lag never drain to zero, even on a healthy consumer?",
        ],
    )

    # ------------------ LAB 7 ------------------
    add_lab_block(doc, num=7,
        title="Structured Logging with Loki",
        goal="Treat logs like metrics — searchable, filterable, correlated with dashboards.",
        context=("Loki is Grafana's log database. It indexes by LABEL "
                 "(cheap) rather than by content (expensive like "
                 "Elasticsearch). Promtail ships every container's "
                 "stdout/stderr into Loki with no agent install."),
        steps=[
            ("step", "1. In Grafana, open the Explore view from the left sidebar."),
            ("step", "2. Switch the data source dropdown (top-left) to 'Loki'."),

            ("step", "3. Run a basic LogQL query to see broker logs."),
            ("code", "{container=\"kafka1\"}"),
            ("why",  "selector form {label=value} streams logs from the kafka1 container."),

            ("step", "4. Filter for errors across all brokers."),
            ("code", "{container=~\"kafka.*\"} |= \"ERROR\""),
            ("why",  "regex selector matches kafka1/2/3 and the '|= \"ERROR\"' pipe keeps only lines containing the word ERROR."),

            ("step", "5. Parse a structured JSON log and filter by a field."),
            ("code", "{container=\"taxi-dashboard\"} | json"),
            ("why",  "streams dashboard JSON logs and parses each line so individual fields become searchable in the bottom panel."),

            ("step", "6. Correlate with metrics — pick the time range of a recent surge spike from your Grafana dashboard and apply it here."),
        ],
        observe=[
            "Logs from every container ship to Loki via Promtail without installing an agent inside the container.",
            "LogQL syntax mirrors PromQL — same mental model: selector + pipes.",
            "Searching across millions of log lines stays cheap because indexing is label-only.",
        ],
        discuss=[
            "Why is Loki cheaper to run than Elasticsearch for high-volume container logs?",
            "What is the trade-off? (Hint: full-text search is slower.)",
        ],
    )

    # ============================================================
    # PART 3 — ERROR HANDLING & DATA QUALITY
    # ============================================================
    doc.add_page_break()
    add_heading(doc, "Part 3 · Error Handling & Data Quality (Labs 8-10)", level=1)
    add_para(doc,
        "Bad data is inevitable. Quarantine it, configure durable retries, "
        "and enforce declarative expectations on every record.",
        italic=True, color=MUTED)

    # ------------------ LAB 8 ------------------
    add_lab_block(doc, num=8,
        title="Error Handling + Dead Letter Queue",
        goal="Bad data is inevitable. Don't let it crash the pipeline — quarantine it.",
        context=("A poison-pill message (bad JSON, missing required "
                 "field, out-of-range value) will jam a naïve consumer "
                 "forever. The DLQ pattern routes failures to a separate "
                 "topic so the main pipeline keeps moving and the "
                 "failures can be triaged later."),
        tip=("The simulator injects ~1.5 % corrupted records on purpose. "
             "The exact rate lives in Lab_Files/config.json → "
             "bad_record_rate."),
        steps=[
            ("step", "1. Make sure the simulator + consumer from Lab 4 are still running."),
            ("why",  "if you stopped them, restart taxi_simulator.py and taxi_consumer.py before continuing."),

            ("step", "2. Inspect the DLQ — read the last 100 records and group by reason."),
            ("code", "python dlq_tool.py --mode inspect --limit 100"),
            ("why",  "connects to the trips-dlq topic with --from-beginning, reads up to 100 messages and prints counts per _dlq_reason header. Typical output: neg_fare 34, outlier_dist 22, missing_driver_id 18, bad_coord 16, future_ts 10."),

            ("step", "3. Watch the DLQ counter on the live dashboard."),
            ("why",  "it should be slowly ticking up — about 1.5 % of trips fail validation as configured. The right-hand DLQ table shows the most recent bad records with their reason."),

            ("step", "4. Once you've identified the bug downstream (in a real incident), replay the DLQ."),
            ("code", "python dlq_tool.py --mode replay"),
            ("why",  "re-publishes every DLQ record back to the SOURCE topic (taxi-trips), NOT directly to trips-clean. That way validation runs again — replayed-but-still-bad records simply go back into the DLQ."),
        ],
        observe=[
            "The pipeline NEVER crashes — even on garbage input.",
            "trips-dlq has 24 h retention → you have time to triage.",
            "Replay routes back through validation, so 'replay' is safe and idempotent.",
        ],
        discuss=[
            "Why route to trips-dlq instead of just logging and dropping?",
            "What metadata should a DLQ message always carry? (We add _dlq_reason and _dlq_ts.)",
        ],
    )

    # ------------------ LAB 9 ------------------
    add_lab_block(doc, num=9,
        title="Idempotent Producers, Retries, Backoff",
        goal=("Configure a producer for AT-LEAST-ONCE WITH DEDUPLICATION "
              "(effectively exactly-once per partition)."),
        context=("Network blips happen. A producer sends a batch, the "
                 "broker writes it, then the ack is lost. The producer "
                 "retries — result: duplicates in the topic. "
                 "enable.idempotence=true asks the broker to attach a "
                 "producer ID + sequence to every batch and reject "
                 "duplicates. Free safety."),
        steps=[
            ("step", "1. Open taxi_simulator.py and read the producer config block."),
            ("code",
                "{\n"
                "  \"acks\": \"all\",\n"
                "  \"enable.idempotence\": True,\n"
                "  \"compression.type\": \"snappy\",\n"
                "  \"linger.ms\": 50,\n"
                "  \"batch.size\": 32 * 1024,\n"
                "}"),
            ("why",  "these are the four settings that turn a 'cheap producer' into a production-grade one. acks=all + enable.idempotence=True is the standard recipe."),

            ("step", "2. Leave taxi_simulator.py and the rest of the pipeline running. In a NEW terminal, force a broker failure mid-stream."),
            ("code", "docker pause kafka2"),
            ("why",  "stops the kafka2 process without killing the container. The producer immediately sees retries against kafka2 fail, but kafka1 and kafka3 are still in the ISR — writes continue."),

            ("step", "3. Watch the simulator's log output for retry messages."),
            ("why",  "you should see warnings about kafka2 then continued sends. No data is lost because min.insync.replicas=2 is satisfied by the other two brokers."),

            ("step", "4. Resume kafka2."),
            ("code", "docker unpause kafka2"),
            ("why",  "kafka2 catches up the missed messages and rejoins the ISR. You can see this in Kafka UI → Brokers → kafka2 → 'In-Sync Replicas'."),
        ],
        observe=[
            "With acks=all + min.insync.replicas=2 the producer waits until at least 2 brokers ack each write.",
            "With enable.idempotence=true, retries do NOT create duplicates — the broker dedupes by producer ID + sequence.",
            "The simulator pauses briefly during the failover, then resumes — no message loss.",
        ],
        discuss=[
            "What changes if you set acks=1?",
            "When is enable.idempotence=false acceptable?",
        ],
    )

    # ------------------ LAB 10 ------------------
    add_lab_block(doc, num=10,
        title="Expectation-Based Validation on a Streaming Pipeline",
        goal=("Apply declarative data-quality rules (Great Expectations / "
              "Soda / Deequ style) to micro-batches off Kafka."),
        context=("The validator runs each rule on every record, "
                 "increments a per-rule Prometheus counter, and "
                 "quarantines failures to the DLQ with the rule name "
                 "preserved. Rules live in config.json — no code changes "
                 "needed to retune."),
        tip=("Quality rules currently enforced (all in config.json → "
             "'quality' block): trip_id not null · driver_id not null · "
             "fare_amount 0..1000 · distance_miles 0..150 · "
             "passenger_count 1..8 · payment_type in {credit_card, cash, "
             "mobile, corporate} · pickup_lat 40.4..41.0 · pickup_lon "
             "-74.3..-73.5 (NYC bounding box)."),
        steps=[
            ("step", "1. Watch the validator's stdout — it prints batch results every 25 records."),
            ("code",
                "batch=25 score=87.5% passed=7/8\n"
                "FAIL distance_in_range  sample=[9999.0]"),
            ("why",  "'score' is the overall quality score for that micro-batch. Failing rules name themselves and show a sample bad value."),

            ("step", "2. In Grafana, find the 'Data Quality Score' tile on the dashboard."),
            ("why",  "it reads the data_quality_score gauge exported by quality_validator.py."),

            ("step", "3. Open Lab_Files/config.json and tighten one rule:"),
            ("code",
                "// config.json → quality block\n"
                "\"passenger_max\": 4     // was 8"),
            ("why",  "editing config.json lets you change a quality threshold WITHOUT touching code."),

            ("step", "4. Restart the validator so it picks up the new threshold."),
            ("code",
                "# In Terminal 5, Ctrl+C then:\n"
                "python quality_validator.py"),
            ("why",  "watch the score drop — trips with 5-8 passengers (a small fraction) now fail and go to the DLQ."),

            ("step", "5. Look at per-rule counters in Prometheus."),
            ("code", "sum by (expectation) (gx_expectations_failed_total)"),
            ("why",  "returns ONE bar per failing rule — instant diagnosis of WHICH expectation broke."),
        ],
        observe=[
            "Each rule emits its own gx_expectations_passed_total{expectation=\"…\"} and a corresponding _failed_ counter.",
            "Rows that fail any rule are routed to trips-dlq with the rule name preserved as _dlq_reason.",
            "Changing a threshold needs ZERO code edits — only config.json + a restart.",
        ],
        discuss=[
            "Streaming vs batch validation: what changes?",
            "Why an in-house validator instead of Great Expectations here? "
            "(GX adds heavy dependencies that don't always work on newer "
            "Python; the 30-line Expectation class keeps the IDEA without "
            "the install pain.)",
            "When would you reject the WHOLE batch vs. reject individual rows?",
        ],
    )

    # ============================================================
    # PART 4 — STREAM PROCESSING + ALERTING
    # ============================================================
    doc.add_page_break()
    add_heading(doc, "Part 4 · Stream Processing & Alerting (Labs 11-12)", level=1)
    add_para(doc,
        "Implement windowed aggregations both in Python and in ksqlDB, "
        "then route metric breaches to Alertmanager.",
        italic=True, color=MUTED)

    # ------------------ LAB 11 ------------------
    add_lab_block(doc, num=11,
        title="Windowed Aggregations — Surge Pricing in Python",
        goal=("Implement a real Kafka-Streams-style tumbling-window "
              "aggregation in pure Python so the logic is VISIBLE."),
        context=("Kafka Streams in Java hides the windowing machinery. "
                 "In Python you see every line. surge_detector.py rolls "
                 "a 30-second tumbling window per pickup zone, counts "
                 "demand and supply, computes a surge multiplier, and "
                 "publishes it to a compacted topic."),
        steps=[
            ("step", "1. Open Lab_Files/surge_detector.py in your editor and read it top-to-bottom."),
            ("why",  "focus on compute_surge(demand, supply) — that's the business logic. The rest is plumbing."),

            ("step", "2. Make sure it is running (Terminal 4 from Lab 4)."),

            ("step", "3. Watch its stdout."),
            ("code",
                "zone=TIMES_SQUARE  demand=14 supply= 2 surge=2.00x\n"
                "zone=JFK_AIRPORT   demand= 6 supply= 1 surge=1.60x\n"
                "zone=MIDTOWN       demand= 4 supply= 5 surge=1.00x"),
            ("why",  "every 30 s it prints one line per zone with its current surge multiplier."),

            ("step", "4. Confirm the surge-events topic is being written. It is COMPACTED — only the latest value per zone is kept."),
            ("code",
                "docker exec kafka1 kafka-console-consumer \\\n"
                "  --bootstrap-server kafka1:29092 \\\n"
                "  --topic surge-events \\\n"
                "  --from-beginning --max-messages 5 \\\n"
                "  --property print.key=true"),
            ("why",  "each record's key is the zone name; the value is the latest surge multiplier. Compaction garbage-collects older values per key."),

            ("step", "5. Open the live dashboard (http://localhost:5000) and watch the zone circles change colour as the surge multiplier moves."),
        ],
        observe=[
            "compute_surge(demand, supply) is the SAME shape as a Kafka Streams aggregate() call in Java.",
            "TUMBLING = non-overlapping windows. SLIDING/HOPPING = overlapping windows.",
            "Compaction means a downstream consumer can rewind to the start of the topic and still see the LATEST surge per zone.",
        ],
        discuss=[
            "Where would you use a SLIDING window instead of tumbling?",
            "What happens to a zone's surge value at midnight if demand drops to 0? "
            "(The compacted topic keeps the last value forever unless we publish a new one.)",
        ],
    )

    # ------------------ LAB 12 ------------------
    add_lab_block(doc, num=12,
        title="ksqlDB + Alertmanager",
        goal=("Express stream processing as SQL that runs continuously "
              "inside Kafka, then route metric breaches to Alertmanager."),
        context=("ksqlDB lets you write streaming SQL. The CREATE STREAM "
                 "/ CREATE TABLE statements compile down to Kafka "
                 "Streams topologies that run continuously. Useful when "
                 "SQL is enough and you'd rather not write Java/Python."),
        steps=[
            ("step", "1. Part A — ksqlDB. Connect to the ksqlDB CLI."),
            ("code",
                "docker exec -it ksqldb-cli ksql http://ksqldb-server:8088"),
            ("why",  "drops you into the interactive 'ksql>' prompt inside the ksqldb-cli container."),

            ("step", "2. Tell ksqlDB to read every topic from the BEGINNING for this session."),
            ("code", "SET 'auto.offset.reset' = 'earliest';"),
            ("why",  "without this, your queries only see records arriving AFTER you start them. With it, ksqlDB reads what's already on disk too."),

            ("step", "3. Paste the statements from ksql_taxi.sql one at a time. Start with the stream definition."),
            ("code",
                "CREATE STREAM trips_raw (...)\n"
                "  WITH (KAFKA_TOPIC='taxi-trips', VALUE_FORMAT='JSON');"),
            ("why",  "registers an existing Kafka topic as a ksqlDB STREAM so you can query it with SQL."),

            ("step", "4. Inspect what ksqlDB now knows about."),
            ("code", "SHOW STREAMS;\nSHOW TABLES;"),

            ("step", "5. Run a PUSH query — continuous, streams results forever."),
            ("code", "SELECT * FROM surge_trips EMIT CHANGES LIMIT 5;"),
            ("why",  "PUSH queries push new rows to you as they arrive. EMIT CHANGES is required for that."),

            ("step", "6. Run a PULL query — point-in-time lookup."),
            ("code",
                "SELECT pickup_zone, trip_count, revenue\n"
                "FROM revenue_per_zone_per_minute\n"
                "WHERE pickup_zone='TIMES_SQUARE';"),
            ("why",  "PULL queries return the current value for a key. Like a normal SQL SELECT, single result, then done."),

            ("step", "7. Exit the ksqlDB CLI."),
            ("code", "EXIT;"),

            ("step", "8. Part B — Alerting. Open Lab_Files/alerts.yml and read the five rules."),
            ("why",  "rules defined: HighConsumerLag (lag > 1000) · BrokerDown · UnderReplicatedPartitions · HighErrorRate (DLQ rate > 5/s) · SurgeActive (any zone > 1.5× for 30s)."),

            ("step", "9. Trigger HighConsumerLag — kill the consumer for ~2 minutes."),
            ("code", "# In Terminal 2: Ctrl+C taxi_consumer.py and wait."),
            ("why",  "lag will climb past 1000 within seconds. The alert moves from inactive → pending → firing after the 'for:' duration in alerts.yml."),

            ("step", "10. Check the alert status."),
            ("code",
                "http://localhost:9090/alerts          # Prometheus side\n"
                "http://localhost:9093                  # Alertmanager side"),
            ("why",  "Prometheus shows where each alert is in its lifecycle. Alertmanager shows what notifications fired (we use the 'null' receiver here so nothing actually pages)."),

            ("step", "11. Restart the consumer to clear the alert."),
            ("code", "python taxi_consumer.py"),
        ],
        observe=[
            "ksqlDB writes its query outputs to NEW Kafka topics (surge_trips, taxi_zone_revenue_1m, …).",
            "TUMBLING vs HOPPING vs SESSION windows produce visibly different table cardinalities.",
            "Alerts transition: inactive → pending (during the 'for:' interval) → firing.",
        ],
        discuss=[
            "When is ksqlDB the right tool vs. Kafka Streams (Java) or a custom Python consumer?",
            "Why does an alert have a 'for:' duration? Why not fire immediately?",
        ],
    )

    # ============================================================
    # PART 5 — SCALING
    # ============================================================
    doc.add_page_break()
    add_heading(doc, "Part 5 · Scaling (Labs 13-15)", level=1)
    add_para(doc,
        "Scale a consumer group, add a broker to the cluster live, and "
        "measure throughput across producer-tuning combinations — then "
        "kill a broker to prove the cluster survives.",
        italic=True, color=MUTED)

    # ------------------ LAB 13 ------------------
    add_lab_block(doc, num=13,
        title="Horizontal Scaling — Partitions & Consumers",
        goal="Scale a consumer group by adding more instances, then more partitions.",
        context=("ONE partition can be consumed by AT MOST ONE consumer "
                 "in a group. So adding consumers helps only up to the "
                 "partition count. Beyond that you must add partitions "
                 "too."),
        steps=[
            ("step", "1. With ONE consumer running, look at the partition assignment."),
            ("code",
                "docker exec kafka1 kafka-consumer-groups --bootstrap-server kafka1:29092 \\\n"
                "  --describe --group trip-processor-v1"),
            ("why",  "shows which consumer instance owns which partitions. With one consumer, it owns all 6 partitions of taxi-trips."),

            ("step", "2. Open a NEW terminal (activate the venv) and start a SECOND consumer in the same group."),
            ("code",
                "# new terminal, Lab_Files/, venv activated\n"
                "python taxi_consumer.py"),
            ("why",  "same script, same group ID → Kafka triggers a REBALANCE so the two consumers split the work."),

            ("step", "3. Re-describe the group."),
            ("code",
                "docker exec kafka1 kafka-consumer-groups --bootstrap-server kafka1:29092 \\\n"
                "  --describe --group trip-processor-v1"),
            ("why",  "partitions are now split 3+3 across the two consumers."),

            ("step", "4. Start a THIRD consumer (3+3+2 with 6 partitions) and then a FOURTH (one will be IDLE)."),
            ("why",  "with 4 consumers but only 6 partitions, the fourth gets at most one. Adding a fifth would leave one consumer with no partitions at all."),

            ("step", "5. Add partitions on the fly to use the idle consumers."),
            ("code",
                "docker exec kafka1 kafka-topics --bootstrap-server kafka1:29092 \\\n"
                "  --alter --topic taxi-trips --partitions 12"),
            ("why",  "doubles the partition count from 6 to 12. Kafka rebalances and now all 4 consumers get work (3+3+3+3)."),

            ("step", "6. Stop the extra consumers (Ctrl+C in their terminals) when you're done."),
        ],
        observe=[
            "Partition count is the UPPER LIMIT on parallelism inside a single consumer group.",
            "Rebalances appear in consumer logs: 'Revoking partitions' → 'Assigning partitions'.",
            "Adding partitions LIVE works — but it changes which key goes to which partition, breaking ordering guarantees for any future records.",
        ],
        discuss=[
            "Why can you never DECREASE partitions on a topic?",
            "What happens to ordering guarantees if you change the partitioner?",
        ],
    )

    # ------------------ LAB 14 ------------------
    add_lab_block(doc, num=14,
        title="Vertical Scaling — Add a Broker, Rebalance Partitions",
        goal="Grow the cluster from 3 → 4 brokers and move data over.",
        context=("Adding a broker is two operations: bring the new broker "
                 "into the quorum, then ASK Kafka to move some existing "
                 "replicas onto it. Kafka does not auto-rebalance — you "
                 "control when the migration happens."),
        warn=("This lab edits docker-compose.yml. Take a backup copy "
              "first so you can restore the original layout when you're "
              "done."),
        steps=[
            ("step", "1. Open docker-compose.yml and append a 'kafka4' service modelled on kafka3:"),
            ("step", "    • set KAFKA_NODE_ID: 4"),
            ("step", "    • add 4@kafka4:9093 to KAFKA_CONTROLLER_QUORUM_VOTERS on ALL FOUR nodes (kafka1/2/3/4)"),
            ("step", "    • map host port 9096:9096 and set EXTERNAL://0.0.0.0:9096"),
            ("why",  "the new broker has to know the quorum voters, and the existing brokers have to know about it."),

            ("step", "2. Bring up the new broker."),
            ("code", "docker compose up -d kafka4"),
            ("why",  "starts only the kafka4 service. Within ~30 s it joins the KRaft quorum. Verify in Kafka UI → Brokers."),

            ("step", "3. Generate the list of topics to move (Linux/macOS/WSL)."),
            ("code",
                "cat > reassign-topics.json <<'JSON'\n"
                "{\"topics\":[{\"topic\":\"taxi-trips\"}],\"version\":1}\n"
                "JSON"),
            ("step", "    Windows cmd version:"),
            ("code",
                "echo {\"topics\":[{\"topic\":\"taxi-trips\"}],\"version\":1} > reassign-topics.json"),
            ("why",  "this is the INPUT for kafka-reassign-partitions: which topic(s) to consider for the move."),

            ("step", "4. Copy the file into the kafka1 container and generate the reassignment plan."),
            ("code",
                "docker cp reassign-topics.json kafka1:/tmp/reassign-topics.json\n"
                "docker exec kafka1 kafka-reassign-partitions \\\n"
                "  --bootstrap-server kafka1:29092 \\\n"
                "  --topics-to-move-json-file /tmp/reassign-topics.json \\\n"
                "  --broker-list \"1,2,3,4\" --generate"),
            ("why",  "Kafka prints a 'Proposed partition reassignment configuration' JSON block. Copy the entire block — that is your plan."),

            ("step", "5. Save the proposed JSON into a file called reassignment.json on your host, then copy it back into kafka1 and execute it."),
            ("code",
                "docker cp reassignment.json kafka1:/tmp/reassignment.json\n"
                "docker exec kafka1 kafka-reassign-partitions \\\n"
                "  --bootstrap-server kafka1:29092 \\\n"
                "  --reassignment-json-file /tmp/reassignment.json --execute"),
            ("why",  "kicks off the data move. Some partitions on kafka1/2/3 will be copied onto kafka4, then the leadership transfers."),

            ("step", "6. Verify the reassignment finishes."),
            ("code",
                "docker exec kafka1 kafka-reassign-partitions \\\n"
                "  --bootstrap-server kafka1:29092 \\\n"
                "  --reassignment-json-file /tmp/reassignment.json --verify"),
            ("why",  "reports 'Reassignment of partition X is completed' for each affected partition."),
        ],
        observe=[
            "During reassignment, 'Bytes Out' on brokers 1/2/3 spikes — they're shipping data to broker 4.",
            "UnderReplicatedPartitions is briefly > 0, then returns to 0.",
            "Topic counts in Kafka UI show the new replica distribution including kafka4.",
        ],
        discuss=[
            "Why doesn't Kafka rebalance partitions automatically when you add a broker?",
            "What is 'Cruise Control' and how would it automate this?",
        ],
    )

    # ------------------ LAB 15 ------------------
    add_lab_block(doc, num=15,
        title="Throughput Tuning + Broker-Failure Resilience",
        goal=("Measure how producer settings change throughput, then "
              "prove the cluster survives a broker loss with no data "
              "loss."),
        context=("Part A is empirical — change one knob, measure "
                 "throughput, compare. Part B is operational — kill a "
                 "broker and verify the durability settings from Lab 9 "
                 "actually work."),
        steps=[
            ("step", "1. Part A — Throughput tuning. Run the baseline."),
            ("code",
                "python load_test.py --workers 4 --rate 2000 --duration 30"),
            ("why",  "spawns 4 producer workers, each aiming for 2000 msg/s for 30 seconds. Prints the achieved rate per worker at the end."),

            ("step", "2. Edit load_test.py producer config and re-run, varying ONE setting at a time."),
            ("code",
                "linger.ms        : 0  vs 20 vs 100\n"
                "batch.size       : 16K vs 64K vs 256K\n"
                "compression.type : none vs snappy vs lz4 vs zstd\n"
                "acks             : 0 vs 1 vs all"),
            ("why",  "change one setting per run so you can attribute the change in throughput to that setting. Re-run the same load_test.py command after each edit."),

            ("step", "3. For each run capture three numbers."),
            ("step", "    • achieved throughput in msg/s"),
            ("step", "    • taxi_produce_latency_seconds p99 from Prometheus"),
            ("step", "    • per-broker Bytes In from Grafana"),
            ("why",  "these three together tell you 'did the change buy throughput? at what latency cost? and did it stay balanced across brokers?'"),

            ("step", "4. Part B — Broker failure. With the full pipeline still running, KILL broker 2."),
            ("code", "docker kill kafka2"),
            ("why",  "hard-stops the kafka2 container instantly — simulates a power failure. acks=all + min.insync.replicas=2 should keep writes succeeding via kafka1 and kafka3."),

            ("step", "5. Open Kafka UI → Brokers. kafka2 shows 'offline'."),
            ("step", "6. Open Topics → taxi-trips. New leaders are elected for the partitions that had been led by kafka2."),

            ("step", "7. Wait ~60 s. The Prometheus alert 'UnderReplicatedPartitions' fires."),
            ("code", "http://localhost:9090/alerts"),
            ("why",  "confirms the alert pipeline you wired in Lab 12 actually fires on a real failure."),

            ("step", "8. Watch the live dashboard — taxi trips continue without interruption."),

            ("step", "9. Bring kafka2 back."),
            ("code", "docker start kafka2"),
            ("why",  "container restarts. kafka2 catches up the missed messages from the other two brokers' logs, then rejoins the ISR."),

            ("step", "10. Watch UnderReplicatedPartitions drop back to 0."),
        ],
        observe=[
            "linger.ms > 0 + larger batch.size ⇒ higher throughput, slightly higher latency.",
            "compression.type=zstd typically wins on text-heavy payloads.",
            "acks=0 is dangerously fast — data loss on broker failure.",
            "KRaft controllers handle leader election faster than the old Zookeeper-based controller.",
            "No messages lost during the kafka2 kill because min.insync.replicas=2 and acks=all.",
        ],
        discuss=[
            "Where is the bottleneck on your laptop — CPU, network, or disk?",
            "What would happen if you killed TWO brokers simultaneously on this 3-broker cluster?",
        ],
    )

    # ============================================================
    # WRAP-UP
    # ============================================================
    doc.add_page_break()
    add_heading(doc, "Wrap-Up & Capstone Questions", level=1)
    add_para(doc, "After all 15 labs you should be able to answer:")
    for q in [
        "Name three things you would alert on for a Kafka pipeline in production.",
        "Describe the path of a single bad record from producer → DLQ → replay.",
        "Walk through what happens when you add a partition to a topic with 3 active consumers.",
        "Why KRaft over Zookeeper? Give two operational reasons.",
        "How does CDC change your architecture compared to periodic batch sync?",
        "Where would Kubernetes / Minikube actually help here that Docker Compose doesn't?",
    ]:
        p = doc.add_paragraph(q, style="List Number")
        p.paragraph_format.space_after = Pt(2)

    add_heading(doc, "When to Use Kubernetes (the Minikube question)", level=2)
    add_para(doc,
        "You do NOT need Minikube for these labs — Docker Compose handles "
        "multi-broker, scaling consumers and broker-failure scenarios on "
        "a single machine.")
    add_para(doc, "Use Kubernetes / Helm (e.g. the Strimzi operator) when you want:")
    add_bullets(doc, [
        "Self-healing: pods restart automatically on node failure.",
        "Autoscaling: HPA / KEDA scales consumer pods based on Kafka lag.",
        "Multi-node distribution: brokers spread across physical nodes for true HA.",
        "Rolling upgrades: rebalance and restart brokers one at a time without downtime.",
    ])
    add_para(doc,
        "Minikube is good for LEARNING Kubernetes; for production use a "
        "real cluster (EKS / GKE / AKS) with Strimzi or Confluent "
        "Operator.")

    # ============================================================
    # EVIDENCE CHECKLIST
    # ============================================================
    add_heading(doc, "Evidence Checklist (submit these)", level=1)
    add_bullets(doc, [
        "docker compose ps — all 15 containers running.",
        "Kafka UI screenshot showing the topics: taxi-trips, gps-pings, trips-clean, trips-enriched, trips-dlq, surge-events, cdc.public.drivers.",
        "Prometheus Targets page showing kafka-brokers UP.",
        "Grafana 'Taxi Pipeline – Live Overview' dashboard screenshot during a surge event.",
        "DLQ inspection output from python dlq_tool.py --mode inspect --limit 100.",
        "Consumer-group lag before / during / after the consumer restart from Lab 6.",
        "Prometheus alert 'HighConsumerLag' or 'UnderReplicatedPartitions' in the 'firing' state.",
        "Throughput numbers from at least 4 producer-config variants in Lab 15 Part A.",
        "Completed RCA note for the consumer-outage scenario.",
    ])

    # ============================================================
    # CLEANUP
    # ============================================================
    add_heading(doc, "Cleanup", level=1)
    add_para(doc, "When you're completely done with the day, stop and "
                  "remove every container plus their data volumes.")
    add_code_block(doc,
        "cd Lab_Files\n"
        "docker compose down -v\n"
        "deactivate                # leave the Python venv")
    add_why(doc,
        "'docker compose down -v' stops every container and DELETES the "
        "named volumes that hold Kafka log segments and Postgres data. "
        "The next 'docker compose up -d' will start from a clean slate "
        "— you'll need to re-run setup_topics.py, db_seeder.py and "
        "register_connector.py (or just run bootstrap.sh again).")

    # ------------------------------------------------------------
    doc.save(OUT)
    print(f"wrote {OUT}  ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build()
